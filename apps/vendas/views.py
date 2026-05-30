import io
import json
from decimal import Decimal
from datetime import date

from apps.core.tenant import tenant_get_or_404
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from apps.accounts.decorators import admin_required
from apps.core.exportacao import exportar_csv
from apps.core.audit import registrar as audit
from django.http import JsonResponse, HttpResponse
from django.db.models import Sum
from django.contrib import messages

from .models import Proposta, ItemProposta, Lead
from apps.cadastros.models import Cliente

def _empresa(request):
    """Retorna a empresa do usuário ou None para superadmin."""
    return getattr(request, 'empresa', None)


def _qs_empresa(qs, request):
    """
    Aplica filtro de empresa ao queryset.
    Se empresa for None (superadmin), retorna o queryset sem filtro.
    """
    empresa = _empresa(request)
    if empresa is None:
        return qs
    return qs.filter(empresa=empresa)




# ─── Helpers ──────────────────────────────────────────────

def _to_dec(v):
    try:
        return Decimal(str(v or 0))
    except Exception:
        return Decimal('0')


def _to_date(v):
    if not v:
        return None
    try:
        return date.fromisoformat(str(v)[:10])
    except Exception:
        return None


def _proposta_to_dict(p):
    """Serializa uma Proposta para dict (uso na lista e no detalhe)."""
    return {
        'id': p.id,
        'codigo': p.codigo,
        'titulo': p.titulo,
        'cliente_id': p.cliente_id,
        'lead_id': p.lead_id,
        'cliente_nome': p.get_cliente_nome(),
        'cliente_tipo': p.get_cliente_tipo(),
        'projeto_nome': p.projeto_nome,
        'data_emissao': p.data_emissao.isoformat() if p.data_emissao else '',
        'data_validade': p.data_validade.isoformat() if p.data_validade else '',
        'status': p.status,
        'valor_total': float(p.valor_total),
        'condicoes_pagamento': p.condicoes_pagamento,
        'prazo_execucao': p.prazo_execucao,
        'observacoes': p.observacoes,
        'notas_tecnicas': p.notas_tecnicas,
        'tem_financeiro': bool(p.transacao_financeiro_ref),
        'projeto_ref_id': p.projeto_ref_id,
        'dados_orcamento': p.dados_orcamento or {},
        'itens': [
            {
                'descricao': it.descricao,
                'unidade': it.unidade,
                'quantidade': float(it.quantidade),
                'preco_unitario': float(it.preco_unitario),
                'preco_total': float(it.preco_total),
            }
            for it in p.itens.all()
        ],
    }


def _criar_projeto_a_partir_de_proposta(proposta):
    """
    Cria automaticamente um Projeto quando a proposta e aprovada.
    Passa os dados financeiros do orcamento para o JSONField dados do Projeto.
    """
    try:
        from apps.projetos.models import Projeto

        resultado = proposta.dados_orcamento.get('resultado', {})
        nome_projeto = proposta.projeto_nome or proposta.titulo

        finances = []
        if resultado.get('valor_venda'):
            finances.append({
                'tipo': 'receita',
                'descricao': 'Valor de Venda (Proposta)',
                'valor': resultado['valor_venda'],
            })
        if resultado.get('custo_mob'):
            finances.append({
                'tipo': 'despesa',
                'descricao': 'Mao de Obra',
                'valor': resultado['custo_mob'],
            })
        if resultado.get('custo_direto'):
            finances.append({
                'tipo': 'despesa',
                'descricao': 'Custos Diretos',
                'valor': resultado['custo_direto'],
            })
        for cp in resultado.get('custos_percentuais_calculados', []):
            finances.append({
                'tipo': 'despesa',
                'descricao': cp.get('nome', 'Custo Percentual'),
                'valor': cp.get('valor', 0),
            })

        dados_iniciais = {
            'tap': {
                'nome': nome_projeto,
                'status': 'planejamento',
                'dataInicio': str(proposta.data_emissao) if proposta.data_emissao else '',
                'dataConclusao': str(proposta.data_validade) if proposta.data_validade else '',
                'gerente': '',
                'patrocinador': proposta.get_cliente_nome(),
                'objetivo': proposta.titulo,
                'escopo': proposta.notas_tecnicas or '',
                'premissas': proposta.condicoes_pagamento or '',
                'requisitos': '',
                'prazo': proposta.prazo_execucao or '',
                'proposta_codigo': proposta.codigo,
                'valor_contrato': resultado.get('valor_venda', float(proposta.valor_total)),
                'margem_prevista': resultado.get('margem_lucro', 0),
                'margem_percentual': resultado.get('margem_percentual', 0),
                'alteracoesEscopo': [],
            },
            'eapTasks': [],
            'finances': finances,
            'kpis': [],
            'risks': [],
            'lessons': [],
            'close': {},
            'actionPlan': [],
        }

        projeto = Projeto.objects.create(
            nome=nome_projeto,
            status='planejamento',
            data_inicio=proposta.data_emissao,
            data_conclusao=proposta.data_validade,
            gerente='',
            patrocinador=proposta.get_cliente_nome(),
            dados=dados_iniciais,
        )

        proposta.projeto_ref_id = projeto.pk
        proposta.save(update_fields=['projeto_ref_id'])
        return projeto

    except Exception as e:
        return None


# ─── Lista principal ────────────────────────────────────────────

@admin_required
def lista(request):
    """Lista de propostas e leads com KPIs. POST via JSON para CRUD de leads e delete de proposta."""
    if request.method == 'POST':
        data = json.loads(request.body)
        action = data.get('action')

        if action == 'save_lead':
            rid = data.get('id')
            obj = tenant_get_or_404(Lead, request, pk=int(rid)) if rid else Lead()
            obj.nome = data.get('nome', '').strip()
            obj.empresa = data.get('empresa', '').strip()
            obj.contato = data.get('contato', '').strip()
            obj.email = data.get('email', '').strip()
            obj.estagio = data.get('estagio', 'prospeccao')
            obj.valor_estimado = _to_dec(data.get('valor_estimado', 0))
            obj.observacoes = data.get('observacoes', '').strip()
            if obj.pk is None and _empresa(request):

                obj.empresa = _empresa(request)

            obj.save()
            return JsonResponse({'ok': True, 'id': obj.id})

        elif action == 'delete_lead':
            _qs_empresa(Lead.objects, request).filter(id=data.get('id')).delete()
            return JsonResponse({'ok': True})

        elif action == 'delete_proposta':
            _qs_empresa(Proposta.objects, request).filter(id=data.get('id')).delete()
            return JsonResponse({'ok': True})

        elif action == 'gerar_financeiro':
            prop = tenant_get_or_404(Proposta, request, pk=data.get('id'))
            ref = f"PROP:{prop.id}"
            try:
                from apps.financeiro.models import Transacao
                if _qs_empresa(Transacao.objects, request).filter(referencia=ref).exists():
                    return JsonResponse({'ok': False, 'msg': 'Ja existe lancamento para esta proposta.'})
                t = Transacao.objects.create(
                    descricao=f"Proposta {prop.codigo} — {prop.titulo}",
                    tipo='entrada',
                    valor=prop.valor_total,
                    data_competencia=prop.data_emissao,
                    data_vencimento=prop.data_validade,
                    status='pendente',
                    cliente_id=prop.cliente_id,
                    referencia=ref,
                    observacoes=f"Gerado automaticamente da proposta {prop.codigo}",
                )
                prop.transacao_financeiro_ref = ref
                prop.save(update_fields=['transacao_financeiro_ref'])
                return JsonResponse({'ok': True, 'msg': f'Conta a receber criada (ID {t.id})'})
            except Exception as e:
                return JsonResponse({'ok': False, 'msg': str(e)})

    propostas = Proposta.objects.select_related('cliente', 'lead').prefetch_related('itens')
    propostas_data = [_proposta_to_dict(p) for p in propostas]

    leads_data = list(_qs_empresa(Lead.objects, request).filter().values(
        'id', 'nome', 'empresa', 'contato', 'email', 'estagio', 'valor_estimado', 'observacoes'
    ))
    for l in leads_data:
        l['valor_estimado'] = float(l['valor_estimado'])

    total_valor = Proposta.objects.aggregate(s=Sum('valor_total'))['s'] or 0
    aprovadas = _qs_empresa(Proposta.objects, request).filter(status='aprovada').count()
    pipeline_valor = _qs_empresa(Lead.objects, request).filter(estagio__in=['qualificacao', 'proposta', 'negociacao']
    ).aggregate(s=Sum('valor_estimado'))['s'] or 0

    ctx = {
        'propostas_json': json.dumps(propostas_data),
        'leads_json': json.dumps(leads_data, default=str),
        'total_propostas': Proposta.objects.count(),
        'total_valor': float(total_valor),
        'aprovadas': aprovadas,
        'pipeline_valor': float(pipeline_valor),
    }
    return render(request, 'vendas/lista.html', ctx)


# ─── Detalhe / Nova Proposta ─────────────────────────────────────────

@admin_required
def proposta_nova(request):
    """Cria uma nova proposta e redireciona para a pagina de detalhe."""
    from datetime import date as dt
    p = Proposta.objects.create(
        codigo='',
        titulo='Nova Proposta',
        data_emissao=dt.today(),
        status='rascunho',
    )
    return redirect('vendas:proposta_detalhe', pk=p.pk)


@admin_required
def proposta_detalhe(request, pk):
    """
    Pagina dedicada da proposta com duas abas:
    - Aba 1: Calculo do Orcamento (MOB, custos diretos, percentuais, resultado, graficos)
    - Aba 2: Proposta (itens + exportar Word)

    POST JSON actions:
    - save_cabecalho
    - save_orcamento
    - save_itens
    - mudar_status
    """
    proposta = tenant_get_or_404(Proposta, request, pk=pk)

    if request.method == 'POST':
        data = json.loads(request.body)
        action = data.get('action')

        if action == 'save_cabecalho':
            proposta.codigo = data.get('codigo', '').strip()
            proposta.titulo = data.get('titulo', '').strip()
            proposta.projeto_nome = data.get('projeto_nome', '').strip()
            proposta.data_emissao = _to_date(data.get('data_emissao')) or date.today()
            proposta.data_validade = _to_date(data.get('data_validade'))
            proposta.condicoes_pagamento = data.get('condicoes_pagamento', '').strip()
            proposta.prazo_execucao = data.get('prazo_execucao', '').strip()
            proposta.observacoes = data.get('observacoes', '').strip()
            proposta.notas_tecnicas = data.get('notas_tecnicas', '').strip()

            tipo_origem = data.get('origem_tipo', '')
            origem_id = data.get('origem_id')
            if tipo_origem == 'cliente':
                proposta.cliente_id = int(origem_id) if origem_id else None
                proposta.lead_id = None
            elif tipo_origem == 'lead':
                proposta.lead_id = int(origem_id) if origem_id else None
                proposta.cliente_id = None
            else:
                proposta.cliente_id = None
                proposta.lead_id = None

            proposta.save()
            return JsonResponse({'ok': True, 'cliente_nome': proposta.get_cliente_nome()})

        elif action == 'save_orcamento':
            proposta.dados_orcamento = data.get('dados_orcamento', {})
            resultado = proposta.dados_orcamento.get('resultado', {})
            valor_venda = resultado.get('valor_venda', 0)
            if valor_venda:
                proposta.valor_total = Decimal(str(valor_venda))
            proposta.save(update_fields=['dados_orcamento', 'valor_total'])
            return JsonResponse({'ok': True})

        elif action == 'save_itens':
            itens = data.get('itens', [])
            proposta.itens.all().delete()
            total = Decimal('0')
            for i, it in enumerate(itens):
                qty = _to_dec(it.get('quantidade', 1))
                preco = _to_dec(it.get('preco_unitario', 0))
                subtotal = qty * preco
                total += subtotal
                ItemProposta.objects.create(
                    proposta=proposta,
                    descricao=it.get('descricao', ''),
                    unidade=it.get('unidade', ''),
                    quantidade=qty,
                    preco_unitario=preco,
                    preco_total=subtotal,
                    ordem=i,
                )
            if not proposta.dados_orcamento.get('resultado', {}).get('valor_venda'):
                proposta.valor_total = total
                proposta.save(update_fields=['valor_total'])
            return JsonResponse({'ok': True, 'valor_total': float(proposta.valor_total)})

        elif action == 'mudar_status':
            novo_status = data.get('status', '')
            status_validos = [s[0] for s in Proposta.STATUS_CHOICES]
            if novo_status not in status_validos:
                return JsonResponse({'ok': False, 'msg': 'Status invalido.'})

            proposta.status = novo_status
            proposta.save(update_fields=['status'])

            projeto_id = None
            msg_extra = ''
            if novo_status == 'aprovada' and not proposta.projeto_ref_id:
                projeto = _criar_projeto_a_partir_de_proposta(proposta)
                if projeto:
                    projeto_id = projeto.pk
                    msg_extra = f' Projeto #{projeto.pk} criado em Gestao de Projetos.'

            return JsonResponse({
                'ok': True,
                'status': proposta.status,
                'projeto_id': projeto_id,
                'msg': f'Status atualizado para "{proposta.get_status_display()}".{msg_extra}',
            })

    clientes = list(_qs_empresa(Cliente.objects, request).filter(ativo=True).values('id', 'nome'))
    leads = list(_qs_empresa(Lead.objects, request).filter().values('id', 'nome', 'empresa'))
    proposta_data = _proposta_to_dict(proposta)

    ctx = {
        'proposta': proposta,
        'proposta_json': json.dumps(proposta_data),
        'clientes_json': json.dumps(clientes),
        'leads_json': json.dumps(leads),
        'status_choices': Proposta.STATUS_CHOICES,
    }
    return render(request, 'vendas/proposta_detalhe.html', ctx)


# ─── Exportar Word ────────────────────────────────────────────────────

@admin_required
def exportar_word(request, pk):
    """
    Gera um arquivo .docx com a proposta formatada para envio ao cliente.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io

    proposta = tenant_get_or_404(Proposta, request, pk=pk)
    itens = proposta.itens.all()

    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def fmt_br(valor):
        try:
            v = float(valor)
            partes = f"{v:,.2f}".split('.')
            inteiro = partes[0].replace(',', '.')
            return f"R$ {inteiro},{partes[1]}"
        except Exception:
            return str(valor)

    AZUL_ESCURO = (30, 58, 138)
    AZUL_MEDIO  = (30, 64, 175)
    AZUL_CLARO  = 'DBEAFE'
    AZUL_LINHA  = 'EFF6FF'
    CINZA_TEXTO = (55, 65, 81)

    # Cabeçalho
    p_empresa = doc.add_paragraph()
    p_empresa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_empresa.add_run('BK ENGENHARIA E TECNOLOGIA')
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(*AZUL_ESCURO)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_titulo.add_run('PROPOSTA TÉCNICA E COMERCIAL')
    r2.bold = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(*AZUL_MEDIO)

    p_cod = doc.add_paragraph()
    p_cod.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_cod.add_run(f'Proposta Nº {proposta.codigo}')
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(*CINZA_TEXTO)
    r3.italic = True

    doc.add_paragraph()

    # Tabela de informações
    tbl_info = doc.add_table(rows=0, cols=2)
    tbl_info.style = 'Table Grid'
    tbl_info.columns[0].width = Cm(5)
    tbl_info.columns[1].width = Cm(11)

    campos = [
        ('Título',              proposta.titulo or '—'),
        ('Cliente',             proposta.get_cliente_nome() or '—'),
        ('Projeto',             proposta.projeto_nome or '—'),
        ('Data de Emissão',     proposta.data_emissao.strftime('%d/%m/%Y') if proposta.data_emissao else '—'),
        ('Válida Até',          proposta.data_validade.strftime('%d/%m/%Y') if proposta.data_validade else '—'),
        ('Prazo de Execução',   proposta.prazo_execucao or '—'),
        ('Condições de Pagto.', proposta.condicoes_pagamento or '—'),
    ]

    for label, valor in campos:
        row = tbl_info.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        set_cell_bg(c0, AZUL_CLARO)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = RGBColor(*AZUL_ESCURO)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(valor)
        r1.font.size = Pt(10)

    doc.add_paragraph()

    # Escopo do serviço
    if proposta.notas_tecnicas:
        p_sec = doc.add_paragraph()
        r_sec = p_sec.add_run('ESCOPO DO SERVIÇO')
        r_sec.bold = True
        r_sec.font.size = Pt(11)
        r_sec.font.color.rgb = RGBColor(*AZUL_ESCURO)
        p_sec.paragraph_format.space_after = Pt(4)
        for linha in proposta.notas_tecnicas.split('\n'):
            p = doc.add_paragraph(linha.strip())
            p.paragraph_format.left_indent = Cm(0.5)
            for run in p.runs:
                run.font.size = Pt(10)
        doc.add_paragraph()

    # Tabela de itens
    p_sec2 = doc.add_paragraph()
    r_sec2 = p_sec2.add_run('ITENS DA PROPOSTA')
    r_sec2.bold = True
    r_sec2.font.size = Pt(11)
    r_sec2.font.color.rgb = RGBColor(*AZUL_ESCURO)
    p_sec2.paragraph_format.space_after = Pt(4)

    col_widths = [Cm(1.0), Cm(7.5), Cm(1.5), Cm(1.8), Cm(2.5), Cm(2.7)]
    headers = ['#', 'Descrição', 'Unid.', 'Qtd.', 'Preço Unit.', 'Total']

    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = 'Table Grid'
    for i, w in enumerate(col_widths):
        tbl.columns[i].width = w

    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '1E3A8A')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)

    total_geral = 0
    for idx_item, item in enumerate(itens):
        bg = AZUL_LINHA if idx_item % 2 == 1 else 'FFFFFF'
        row = tbl.add_row()
        vals = [
            str(idx_item + 1),
            item.descricao or '',
            item.unidade or '',
            str(item.quantidade),
            fmt_br(item.preco_unitario),
            fmt_br(item.preco_total),
        ]
        aligns = [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
        ]
        for i, (v, al) in enumerate(zip(vals, aligns)):
            cell = row.cells[i]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = al
            r = p.add_run(v)
            r.font.size = Pt(10)
        try:
            total_geral += float(item.preco_total)
        except Exception:
            pass

    total_row = tbl.add_row()
    for i in range(6):
        set_cell_bg(total_row.cells[i], AZUL_CLARO)
    total_row.cells[0].merge(total_row.cells[4])
    p_lbl = total_row.cells[0].paragraphs[0]
    p_lbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_lbl = p_lbl.add_run('VALOR TOTAL DA PROPOSTA')
    r_lbl.bold = True
    r_lbl.font.size = Pt(10)
    r_lbl.font.color.rgb = RGBColor(*AZUL_ESCURO)
    p_tot = total_row.cells[5].paragraphs[0]
    p_tot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_tot = p_tot.add_run(fmt_br(total_geral))
    r_tot.bold = True
    r_tot.font.size = Pt(10)
    r_tot.font.color.rgb = RGBColor(*AZUL_ESCURO)

    doc.add_paragraph()

    # Observações
    if proposta.observacoes:
        p_obs_hd = doc.add_paragraph()
        r_obs_hd = p_obs_hd.add_run('OBSERVAÇÕES')
        r_obs_hd.bold = True
        r_obs_hd.font.size = Pt(11)
        r_obs_hd.font.color.rgb = RGBColor(*AZUL_ESCURO)
        p_obs_hd.paragraph_format.space_after = Pt(4)
        for linha in proposta.observacoes.split('\n'):
            p = doc.add_paragraph(linha.strip())
            p.paragraph_format.left_indent = Cm(0.5)
            for run in p.runs:
                run.font.size = Pt(10)
        doc.add_paragraph()

    # Validade
    if proposta.data_validade:
        p_val = doc.add_paragraph()
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_val = p_val.add_run(
            f'Esta proposta é válida até {proposta.data_validade.strftime("%d/%m/%Y")}.'
        )
        r_val.italic = True
        r_val.font.size = Pt(10)
        r_val.font.color.rgb = RGBColor(180, 83, 9)
        doc.add_paragraph()

    # Assinatura
    p_linha = doc.add_paragraph()
    p_linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_linha.add_run('_' * 40)

    p_nome = doc.add_paragraph()
    p_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_nome = p_nome.add_run('BK ENGENHARIA E TECNOLOGIA')
    r_nome.bold = True
    r_nome.font.size = Pt(10)
    r_nome.font.color.rgb = RGBColor(*AZUL_ESCURO)

    p_rodape = doc.add_paragraph()
    p_rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_rod = p_rodape.add_run('Engenharia com Excelência | contato@bk-engenharia.com')
    r_rod.font.size = Pt(9)
    r_rod.font.color.rgb = RGBColor(*CINZA_TEXTO)
    r_rod.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    nome_arquivo = f"Proposta_{proposta.codigo}.docx"
    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{nome_arquivo}"'
    return response



@admin_required
def exportar_propostas(request):
    empresa = _empresa(request)
    qs = Proposta.objects.filter(empresa=empresa).values('id', 'titulo', 'cliente__nome', 'status', 'valor_total', 'data_criacao')
    rows = [list(r.values()) for r in qs]
    return exportar_csv('propostas.csv', ['ID', 'Título', 'Cliente', 'Status', 'Valor Total', 'Data'], rows)
