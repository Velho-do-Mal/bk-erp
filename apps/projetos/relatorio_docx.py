"""
Geração do Relatório Geral do projeto em Word (.docx), pra apresentação a
patrocinador/gestor. Mesma fonte de dados do relatório HTML (projeto.dados:
tap, eapTasks, finances, risks, kpis) + os documentos reais de Controle de
Projetos (model DocumentoControle).

Mantém deliberadamente a mesma paleta/lógica de cores e o mesmo cálculo de
datas da EAP (calcularDatas() em detalhe.html), pra que o relatório impresso
bata com o que o usuário vê na tela.
"""
import io
import textwrap
from datetime import date, datetime, timedelta

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.lines import Line2D

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta (mesma da tela: tabela EAP / Gantt / organograma) ────────────
STATUS_COLOR = {
    'concluido': '#10B981',
    'em-andamento': '#3B82F6',
    'atrasado': '#EF4444',
    'nao-iniciado': '#64748B',
}
STATUS_LABEL = {
    'concluido': 'Concluído',
    'em-andamento': 'Em Andamento',
    'atrasado': 'Atrasado',
    'nao-iniciado': 'Não Iniciado',
}
DOC_STATUS = {
    'nao_iniciado': {'l': 'Não iniciado', 'c': '#64748B'},
    'em_andamento': {'l': 'Em andamento', 'c': '#3B82F6'},
    'em_analise': {'l': 'Em análise', 'c': '#F59E0B'},
    'em_revisao': {'l': 'Em revisão', 'c': '#8B5CF6'},
    'concluido': {'l': 'Aprovado', 'c': '#10B981'},
    'cancelado': {'l': 'Cancelado', 'c': '#EF4444'},
}
RISCO_CORES = {'Crítico': '#EF4444', 'Alto': '#F97316', 'Médio': '#F59E0B', 'Baixo': '#10B981'}
PROB_LABELS = ['Muito Baixa', 'Baixa', 'Média', 'Alta', 'Muito Alta']
IMP_LABELS = ['Muito Baixo', 'Baixo', 'Médio', 'Alto', 'Muito Alto']

FONT_NAME = 'Calibri'
TITLE_RGB = RGBColor(0x1F, 0x38, 0x64)      # #1F3864
HEADER_HEX = '2E5395'                        # fundo dos cabeçalhos de tabela
HEADER_RGB = RGBColor(0x2E, 0x53, 0x95)
ZEBRA_HEX = 'F2F2F2'
MUTED_RGB = RGBColor(0x64, 0x74, 0x8B)


def _status_key(status):
    return status if status in STATUS_COLOR else 'nao-iniciado'


def _hex_to_rgb01(hexcolor):
    h = hexcolor.lstrip('#')
    return tuple(int(h[i:i + 2], 16) / 255 for i in (0, 2, 4))


def _fmt_brl(v):
    try:
        v = float(v or 0)
    except (TypeError, ValueError):
        v = 0.0
    s = f"{v:,.2f}"
    s = s.replace(',', '_').replace('.', ',').replace('_', '.')
    return f"R$ {s}"


def _fmt_data(d):
    if not d:
        return '—'
    if isinstance(d, (date, datetime)):
        return d.strftime('%d/%m/%Y')
    try:
        return datetime.fromisoformat(str(d)[:10]).strftime('%d/%m/%Y')
    except ValueError:
        return str(d)


def _classe_risco(prob, imp):
    try:
        v = int(prob or 0) * int(imp or 0)
    except (TypeError, ValueError):
        v = 0
    if v >= 15:
        return 'Crítico'
    if v >= 9:
        return 'Alto'
    if v >= 4:
        return 'Médio'
    return 'Baixo'


# ── Datas da EAP (replica calcularDatas() de detalhe.html) ──────────────
def _calc_datas(tap_inicio, eap_tasks):
    inicio = tap_inicio or date.today()
    mapa = {str(t.get('codigo', '')).strip(): t for t in eap_tasks}
    memo = {}

    def get_inicio(cod, visitados=None):
        visitados = visitados or set()
        if cod in visitados:
            return 0
        visitados = visitados | {cod}
        t = mapa.get(cod)
        if not t:
            return 0
        preds = [p.strip() for p in str(t.get('predecessoras', '') or '').split(',') if p.strip()]
        max_fim = 0
        for p in preds:
            pm = mapa.get(p)
            if pm:
                p_inicio = get_inicio(p, visitados)
                p_fim = p_inicio + (int(pm.get('duracao') or 0))
                max_fim = max(max_fim, p_fim)
        return max_fim

    out = []
    for t in eap_tasks:
        cod = str(t.get('codigo', '')).strip()
        offset = get_inicio(cod)
        dur = int(t.get('duracao') or 1)
        s = inicio + timedelta(days=offset)
        e = s + timedelta(days=dur)
        nt = dict(t)
        nt['startDate'] = s
        nt['endDate'] = e
        out.append(nt)
    return out


def _compute_summary_set(eap_tasks):
    is_summary = set()
    for t in eap_tasks:
        c = str(t.get('codigo', '') or '').strip()
        parts = c.split('.')
        for i in range(1, len(parts)):
            is_summary.add('.'.join(parts[:i]))
    return is_summary


# ── Gráficos (matplotlib, devolvem PNG em memória) ───────────────────────
def _fig_to_png(fig, dpi=170):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf


def chart_gantt(tasks, is_summary):
    if not tasks:
        return None
    n = len(tasks)
    fig_h = max(2.6, 0.42 * n + 1.0)
    fig, ax = plt.subplots(figsize=(10.2, fig_h))
    ax.set_facecolor('#f8fafc')
    fig.patch.set_facecolor('white')

    labels = []
    for i, t in enumerate(tasks):
        cod = str(t.get('codigo', '') or '').strip()
        nome = t.get('nome', '') or ''
        lbl = f"{cod} — {nome}"
        if len(lbl) > 42:
            lbl = lbl[:41].rstrip() + '…'
        labels.append(lbl)

        s, e = t['startDate'], t['endDate']
        # Duração "real" (do cadastro), não a diferença de datas calculada:
        # _calc_datas() usa duracao-ou-1 (mesma regra do calcularDatas() em
        # JS) só pra nunca ter uma barra de largura zero, então uma tarefa
        # de duração 0 (marco) sempre chega aqui com 1 dia de startDate a
        # endDate — o teste de marco tem que olhar o campo original.
        dur_raw = int(t.get('duracao') or 0)
        dur = (e - s).days
        is_sum = cod in is_summary
        is_marco = dur_raw <= 0
        y = n - 1 - i  # linha 0 = topo

        if is_marco:
            ax.scatter([s], [y], marker='D', s=90,
                       color=STATUS_COLOR[_status_key(t.get('status'))],
                       edgecolor='#0F172A', linewidth=1.1, zorder=5)
            continue

        cor = '#1E293B' if is_sum else STATUS_COLOR[_status_key(t.get('status'))]
        h = 0.32 if is_sum else 0.56
        # trilho de fundo (translúcido)
        ax.barh(y, dur, left=s, height=h, color=cor, alpha=0.22,
                edgecolor=cor, linewidth=1.1, zorder=2)
        # preenchimento sólido = % concluído
        pct = max(0.0, min(100.0, float(t.get('percentualConcluido') or 0))) / 100.0
        if pct > 0:
            ax.barh(y, dur * pct, left=s, height=h, color=cor, zorder=3)
        ax.text(e + timedelta(days=max(1, dur * 0.02)), y, f"{t.get('duracao', 0)}d",
                va='center', ha='left', fontsize=8, color='#334155')

    hoje = date.today()
    ax.axvline(hoje, color='#F59E0B', linestyle=':', linewidth=1.6, zorder=4)
    ax.text(hoje, n - 0.3, ' Hoje', color='#B45309', fontsize=8, fontweight='bold', va='bottom')

    ax.set_yticks(range(n - 1, -1, -1))
    ax.set_yticks(list(range(n)))
    ax.set_yticklabels(list(reversed(labels)), fontsize=8.5, color='#1e293b')
    ax.set_ylim(-0.7, n - 0.3)
    ax.grid(axis='x', color='#e2e8f0', linewidth=0.7)
    ax.grid(axis='y', visible=False)
    for spine in ('top', 'right', 'left'):
        ax.spines[spine].set_visible(False)
    ax.spines['bottom'].set_color('#cbd5e1')
    ax.xaxis_date()
    fig.autofmt_xdate(rotation=30, ha='right')
    ax.tick_params(axis='x', labelsize=8.5, colors='#475569')

    legend_items = [mpatches.Patch(color=STATUS_COLOR[k], label=STATUS_LABEL[k]) for k in
                     ('nao-iniciado', 'em-andamento', 'concluido', 'atrasado')]
    legend_items.append(mpatches.Patch(color='#1E293B', label='Fase (resumo)'))
    legend_items.append(Line2D([0], [0], marker='D', color='none', markerfacecolor='#8B5CF6',
                                markeredgecolor='#0F172A', markersize=8, label='Marco'))
    ax.legend(handles=legend_items, loc='upper center', bbox_to_anchor=(0.5, 1.0 + 1.6 / fig_h),
              ncol=3, frameon=False, fontsize=8.5)

    return _fig_to_png(fig)


def chart_curva_s(tasks):
    """Curva S (planejado × realizado, % acumulado por semana) — mesma
    lógica de renderCurvaS() em detalhe.html."""
    if not tasks:
        return None
    min_date = min(t['startDate'] for t in tasks)
    max_date = max(t['endDate'] for t in tasks)
    semanas = []
    cur = min_date
    while cur <= max_date:
        semanas.append(cur)
        cur = cur + timedelta(days=7)
    if not semanas:
        return None

    total_dias = sum(int(t.get('duracao') or 0) for t in tasks)
    planejado, realizado = [], []
    acum_plan = acum_real = 0.0
    for sem_fim in semanas:
        sem_inicio = sem_fim - timedelta(days=7)
        dias_plan_sem = dias_real_sem = 0.0
        for t in tasks:
            overlap = max(0, (min(t['endDate'], sem_fim) - max(t['startDate'], sem_inicio)).days)
            dias_plan_sem += overlap
            status_key = _status_key(t.get('status'))
            if status_key == 'concluido':
                dias_real_sem += overlap
            elif status_key == 'em-andamento':
                dias_real_sem += overlap * (float(t.get('percentualConcluido') or 0) / 100)
        acum_plan += (dias_plan_sem / total_dias) * 100 if total_dias else 0
        acum_real += (dias_real_sem / total_dias) * 100 if total_dias else 0
        planejado.append(min(100, acum_plan))
        realizado.append(min(100, acum_real))

    fig, ax = plt.subplots(figsize=(8.6, 3.4))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.plot(semanas, planejado, color='#3B82F6', marker='o', markersize=3.5, linewidth=1.8, label='Planejado')
    ax.plot(semanas, realizado, color='#10B981', marker='o', markersize=3.5, linewidth=1.8,
            linestyle='--', label='Realizado')
    ax.axvline(date.today(), color='#F59E0B', linestyle=':', linewidth=1.4)
    ax.set_ylim(0, 105)
    ax.set_ylabel('% Acumulado', fontsize=9, color='#334155')
    ax.set_title('Curva S — Avanço Físico Planejado × Realizado', fontsize=10.5, color='#1F3864',
                 fontweight='bold', pad=10)
    ax.legend(frameon=False, fontsize=9, loc='upper left')
    ax.spines[['top', 'right']].set_visible(False)
    ax.grid(color='#e2e8f0', linewidth=0.7)
    ax.set_axisbelow(True)
    ax.tick_params(axis='both', labelsize=8.5, colors='#475569')
    fig.autofmt_xdate(rotation=30, ha='right')
    return _fig_to_png(fig)


def chart_status_eap(eap_tasks):
    if not eap_tasks:
        return None
    counts = {'concluido': 0, 'em-andamento': 0, 'atrasado': 0, 'nao-iniciado': 0}
    for t in eap_tasks:
        counts[_status_key(t.get('status'))] += 1
    labels = [STATUS_LABEL[k] for k, v in counts.items() if v > 0]
    values = [v for v in counts.values() if v > 0]
    colors = [STATUS_COLOR[k] for k, v in counts.items() if v > 0]
    if not values:
        return None
    fig, ax = plt.subplots(figsize=(4.6, 3.4))
    fig.patch.set_facecolor('white')
    wedges, _, autotexts = ax.pie(
        values, colors=colors, autopct=lambda p: f'{p:.0f}%' if p > 0 else '',
        startangle=90, pctdistance=0.78,
        wedgeprops=dict(width=0.42, edgecolor='white', linewidth=2),
        textprops=dict(color='white', fontsize=9, fontweight='bold'),
    )
    ax.legend(wedges, labels, loc='center left', bbox_to_anchor=(1.0, 0.5), frameon=False, fontsize=9)
    ax.set_title('Status das Atividades (EAP)', fontsize=10.5, color='#1F3864', fontweight='bold', pad=10)
    ax.axis('equal')
    return _fig_to_png(fig)


def chart_controle_status(docs):
    if not docs:
        return None
    counts = {}
    for d in docs:
        k = d.get('status') or 'nao_iniciado'
        counts[k] = counts.get(k, 0) + 1
    labels = [DOC_STATUS.get(k, {}).get('l', k) for k in counts]
    values = list(counts.values())
    colors = [DOC_STATUS.get(k, {}).get('c', '#94A3B8') for k in counts]
    fig, ax = plt.subplots(figsize=(5.6, 3.2))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    bars = ax.bar(labels, values, color=colors, width=0.55)
    for b, v in zip(bars, values):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.05, str(v), ha='center', fontsize=9, color='#1e293b')
    ax.set_title('Documentos por Status (Controle de Projetos)', fontsize=10.5, color='#1F3864',
                 fontweight='bold', pad=10)
    ax.spines[['top', 'right']].set_visible(False)
    ax.tick_params(axis='x', labelsize=8.5, colors='#334155', rotation=15)
    ax.tick_params(axis='y', labelsize=8.5, colors='#64748B')
    ax.yaxis.set_major_locator(plt.MaxNLocator(integer=True))
    ax.grid(axis='y', color='#e2e8f0', linewidth=0.7)
    ax.set_axisbelow(True)
    return _fig_to_png(fig)


def chart_financeiro(finances):
    if not finances:
        return None
    mes_plan, mes_real = {}, {}
    for f in finances:
        if (f.get('tipo') or '') != 'Saída':
            continue
        v = float(f.get('valor') or 0)
        kp = str(f.get('dataPrevista') or '')[:7]
        if kp:
            mes_plan[kp] = mes_plan.get(kp, 0) + v
        if f.get('realizado') and f.get('dataRealizada'):
            kr = str(f.get('dataRealizada'))[:7]
            mes_real[kr] = mes_real.get(kr, 0) + v
    keys = sorted(set(mes_plan) | set(mes_real))
    if not keys:
        return None
    plan_vals = [mes_plan.get(k, 0) for k in keys]
    real_vals = [mes_real.get(k, 0) for k in keys]
    x = range(len(keys))
    fig, ax = plt.subplots(figsize=(8.6, 3.6))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    w = 0.38
    ax.bar([i - w / 2 for i in x], plan_vals, width=w, color='#3B82F6', label='Planejado')
    ax.bar([i + w / 2 for i in x], real_vals, width=w, color='#10B981', label='Realizado')
    ax.set_xticks(list(x))
    ax.set_xticklabels(keys, fontsize=8.5, rotation=30, ha='right', color='#334155')
    ax.set_title('Fluxo de Caixa Planejado × Realizado (mensal)', fontsize=10.5, color='#1F3864',
                 fontweight='bold', pad=10)
    ax.legend(frameon=False, fontsize=9)
    ax.spines[['top', 'right']].set_visible(False)
    ax.grid(axis='y', color='#e2e8f0', linewidth=0.7)
    ax.set_axisbelow(True)
    ax.tick_params(axis='y', labelsize=8.5, colors='#64748B')
    return _fig_to_png(fig)


def chart_organograma(eap_tasks, projeto_nome):
    """Organograma da EAP como imagem — mesmo algoritmo de layout em
    árvore do renderEapOrg() (JS), colorido por status da atividade."""
    if not eap_tasks:
        return None

    class Node:
        __slots__ = ('code', 'label', 'status', 'children', 'depth', 'x', 'y', 'w')

        def __init__(self, code, label, status):
            self.code, self.label, self.status = code, label, status
            self.children, self.depth, self.x, self.y, self.w = [], 0, 0, 0, 0

    root = Node('', (projeto_nome or 'PROJETO').upper(), None)
    node_map = {'': root}
    for i, t in enumerate(eap_tasks):
        cod = str(t.get('codigo') or f't{i}')
        nome = (t.get('nome') or f'Tarefa {i + 1}').upper()
        node_map[cod] = Node(cod, f"{cod} {nome}", _status_key(t.get('status')))
    for i, t in enumerate(eap_tasks):
        cod = str(t.get('codigo') or f't{i}')
        nd = node_map[cod]
        parts = cod.split('.')
        parent = node_map.get('.'.join(parts[:-1])) if len(parts) > 1 else None
        (parent or root).children.append(nd)

    def set_depth(n, d):
        n.depth = d
        for c in n.children:
            set_depth(c, d + 1)
    set_depth(root, 0)

    BW, BH, HG, VG = 2.7, 0.86, 0.65, 1.25

    def subtree_w(n):
        if not n.children:
            return BW
        return max(BW, sum(subtree_w(c) for c in n.children) + HG * (len(n.children) - 1))

    positioned = []

    def layout(n, x, y):
        n.x, n.y = x, y
        positioned.append(n)
        total = subtree_w(n)
        cx = x - total / 2
        for c in n.children:
            cw = subtree_w(c)
            layout(c, cx + cw / 2, y - (BH + VG))
            cx += cw + HG
    layout(root, subtree_w(root) / 2, 0)

    max_depth = max((n.depth for n in positioned), default=0)
    fig_w = max(6.0, subtree_w(root) + 1.0)
    fig_h = max(2.6, (max_depth + 1) * (BH + VG) + 0.6)
    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')

    ROOT_COLOR = '#1E40AF'
    ACCENT = ['#F59E0B', '#8B5CF6', '#14B8A6', '#F472B6', '#22D3EE']

    for n in positioned:
        for c in n.children:
            midy = n.y - BH / 2 - VG / 2
            ax.plot([n.x, n.x, c.x, c.x], [n.y - BH / 2, midy, midy, c.y + BH / 2],
                    color='#94a3b8', linewidth=1.1, zorder=1)

    for n in positioned:
        color = ROOT_COLOR if n.status is None else STATUS_COLOR[n.status]
        rect = mpatches.FancyBboxPatch((n.x - BW / 2, n.y - BH / 2), BW, BH,
                                        boxstyle='round,pad=0.02,rounding_size=0.06',
                                        linewidth=0, facecolor=color, zorder=2)
        ax.add_patch(rect)
        accent = ACCENT[n.depth % len(ACCENT)]
        ax.add_patch(mpatches.Rectangle((n.x - BW / 2 + 0.08, n.y + BH / 2 - 0.12),
                                         BW - 0.16, 0.06, facecolor=accent, linewidth=0, zorder=3))
        # Quebra em até 3 linhas curtas (largura fixa da caixa não pode
        # estourar por cima do nó vizinho); trunca com reticências além
        # disso — o nome completo continua na tabela EAP.
        wrapped = textwrap.wrap(n.label, width=17, max_lines=3, placeholder='…')
        ax.text(n.x, n.y, '\n'.join(wrapped), ha='center', va='center', color='white',
                fontsize=7.0, fontweight='bold', zorder=4, linespacing=1.3)

    ax.set_xlim(-0.5, subtree_w(root) + 0.5)
    ax.set_ylim(-(max_depth + 1) * (BH + VG) + 0.3, BH)
    ax.axis('off')
    return _fig_to_png(fig, dpi=190)


# ── Helpers python-docx ──────────────────────────────────────────────────
def _shade_cell(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_text(cell, text, *, bold=False, color=None, size=10, align=None, font=FONT_NAME):
    cell.text = ''
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run('' if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_heading(doc, text, *, size=15, color=TITLE_RGB, before=14, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run.font.color.rgb = color
    pb = p.paragraph_format
    pb.keep_with_next = True
    # linha fina abaixo do título
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), HEADER_HEX)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def _add_table(doc, headers, rows, *, col_widths_cm=None, small=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                        size=9 if small else 9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(hdr[i], HEADER_HEX)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for cidx, val in enumerate(row):
            text, kwargs = (val, {}) if not isinstance(val, tuple) else val
            align = kwargs.get('align', WD_ALIGN_PARAGRAPH.CENTER if cidx > 0 else WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell_text(cells[cidx], text, bold=kwargs.get('bold', False),
                            color=kwargs.get('color'), size=8.5 if small else 9.5, align=align)
            if ridx % 2 == 1:
                _shade_cell(cells[cidx], ZEBRA_HEX)
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    return table


def _add_kpi_cards(doc, cards):
    """Cards de indicadores (nº, rótulo, cor) em uma linha de tabela sem bordas."""
    table = doc.add_table(rows=1, cols=len(cards))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, (valor, label, cor) in enumerate(cards):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(str(valor))
        r1.bold = True
        r1.font.size = Pt(18)
        r1.font.name = FONT_NAME
        r1.font.color.rgb = RGBColor.from_string(cor.lstrip('#'))
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.size = Pt(8.5)
        r2.font.name = FONT_NAME
        r2.font.color.rgb = MUTED_RGB
        _shade_cell(cell, 'F8FAFC')
    return table


def _add_image(doc, png_buf, *, width_cm=16.5, caption=None):
    if png_buf is None:
        p = doc.add_paragraph('Sem dados suficientes para gerar este gráfico.')
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = MUTED_RGB
        p.runs[0].font.size = Pt(9.5)
        return
    doc.add_picture(png_buf, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = MUTED_RGB


def _set_default_font(doc):
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)


# ── Montagem do documento ────────────────────────────────────────────────
def gerar_relatorio_docx(projeto):
    dados = projeto.dados or {}
    tap = dados.get('tap', {}) or {}
    eap_tasks = dados.get('eapTasks', []) or []
    finances = dados.get('finances', []) or []
    risks = dados.get('risks', []) or []
    kpis = dados.get('kpis', []) or []

    docs_qs = list(projeto.documentos_controle.all().order_by('id')) if hasattr(projeto, 'documentos_controle') else []
    docs = [{
        'codigo': d.doc_numero, 'atividade': d.doc_nome or d.servico_nome,
        'revisao': d.revisao, 'responsavel': d.responsavel_bk,
        'percentual': d.percentual_concluido, 'status': d.status,
    } for d in docs_qs]

    try:
        tap_inicio = date.fromisoformat(str(tap.get('dataInicio'))[:10]) if tap.get('dataInicio') else (
            projeto.data_inicio or date.today())
    except ValueError:
        tap_inicio = projeto.data_inicio or date.today()

    tasks_calc = _calc_datas(tap_inicio, eap_tasks)
    is_summary = _compute_summary_set(eap_tasks)

    doc = Document()
    _set_default_font(doc)

    # ── Capa ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run('RELATÓRIO GERAL DE PROJETO')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = FONT_NAME
    run.font.color.rgb = MUTED_RGB

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(20)
    run2 = p2.add_run(tap.get('nome') or projeto.nome)
    run2.bold = True
    run2.font.size = Pt(24)
    run2.font.name = FONT_NAME
    run2.font.color.rgb = TITLE_RGB

    info_rows = [
        ('Cliente / Patrocinador', tap.get('patrocinador') or projeto.patrocinador or '—'),
        ('Gerente do Projeto', tap.get('gerente') or projeto.gerente or '—'),
        ('Status', (tap.get('status') or projeto.get_status_display() or '—')),
        ('Início Previsto', _fmt_data(tap.get('dataInicio') or projeto.data_inicio)),
        ('Conclusão Prevista', _fmt_data(tap.get('dataConclusao') or projeto.data_conclusao)),
        ('Emitido em', datetime.now().strftime('%d/%m/%Y às %H:%M')),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, val in info_rows:
        row = t.add_row().cells
        _set_cell_text(row[0], label, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, color=MUTED_RGB)
        _set_cell_text(row[1], val, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        row[0].width = Cm(6.5)
        row[1].width = Cm(8.5)
    doc.add_page_break()

    # ── Indicadores gerais (cards) ──────────────────────────────────────
    _add_heading(doc, '1. Indicadores Gerais do Projeto', size=16)
    concl_eap = sum(1 for t in eap_tasks if _status_key(t.get('status')) == 'concluido')
    atras_eap = sum(1 for t in eap_tasks if _status_key(t.get('status')) == 'atrasado')
    pct_medio = (sum(float(t.get('percentualConcluido') or 0) for t in eap_tasks) / len(eap_tasks)) if eap_tasks else 0
    docs_concl = sum(1 for d in docs if d['status'] in ('concluido',))
    riscos_criticos = sum(1 for r in risks if _classe_risco(r.get('probabilidade'), r.get('impacto')) in ('Crítico', 'Alto'))
    cards = [
        (len(eap_tasks), 'Atividades na EAP', '#1F3864'),
        (f"{pct_medio:.0f}%", '% Médio Concluído', '#3B82F6'),
        (concl_eap, 'Atividades Concluídas', '#10B981'),
        (atras_eap, 'Atividades Atrasadas', '#EF4444'),
        (f"{docs_concl}/{len(docs)}" if docs else '—', 'Docs Aprovados', '#8B5CF6'),
        (riscos_criticos, 'Riscos Críticos/Altos', '#F97316'),
    ]
    _add_kpi_cards(doc, cards)
    doc.add_paragraph()

    if kpis:
        kpi_rows = []
        for k in kpis:
            prev, real = float(k.get('previsto') or 0), float(k.get('realizado') or 0)
            diff = real - prev
            cor = RGBColor(0x10, 0xB9, 0x81) if diff >= 0 else RGBColor(0xEF, 0x44, 0x44)
            kpi_rows.append([
                k.get('nome', ''), k.get('unidade', ''), f"M{k.get('mes', 1)}",
                f"{prev:.2f}", f"{real:.2f}", (f"{diff:+.2f}", {'color': cor, 'bold': True}),
            ])
        _add_table(doc, ['KPI', 'Unidade', 'Mês', 'Previsto', 'Realizado', 'Diferença'], kpi_rows, small=True)
        doc.add_paragraph()

    # ── Organograma ───────────────────────────────────────────────────
    _add_heading(doc, '2. Organograma da EAP', size=16)
    _add_image(doc, chart_organograma(eap_tasks, tap.get('nome') or projeto.nome), width_cm=16.5)

    # ── Tabela EAP ────────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, '3. Estrutura Analítica do Projeto (EAP)', size=16)
    if tasks_calc:
        eap_rows = []
        for t in tasks_calc:
            cod = str(t.get('codigo') or '')
            status_key = _status_key(t.get('status'))
            cor = RGBColor.from_string(STATUS_COLOR[status_key].lstrip('#'))
            eap_rows.append([
                cod, t.get('nome', ''), f"{t.get('duracao', 0)}d",
                _fmt_data(t['startDate']), _fmt_data(t['endDate']),
                t.get('responsavel') or '—',
                f"{int(t.get('percentualConcluido') or 0)}%",
                (STATUS_LABEL[status_key], {'bold': True, 'color': cor}),
            ])
        _add_table(doc, ['Cód.', 'Atividade', 'Dur.', 'Início', 'Conclusão', 'Responsável', '%', 'Status'],
                    eap_rows, small=True)
    else:
        doc.add_paragraph('Nenhuma atividade cadastrada na EAP.')

    # ── Gantt ─────────────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, '4. Diagrama de Gantt', size=16)
    _add_image(doc, chart_gantt(tasks_calc, is_summary), width_cm=16.5)
    _add_image(doc, chart_curva_s(tasks_calc), width_cm=16.5)
    _add_image(doc, chart_status_eap(eap_tasks), width_cm=10)

    # ── Financeiro ────────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, '5. Análise Financeira', size=16)
    mes_plan, mes_real = {}, {}
    for f in finances:
        if (f.get('tipo') or '') != 'Saída':
            continue
        v = float(f.get('valor') or 0)
        kp = str(f.get('dataPrevista') or '')[:7]
        if kp:
            mes_plan[kp] = mes_plan.get(kp, 0) + v
        if f.get('realizado') and f.get('dataRealizada'):
            kr = str(f.get('dataRealizada'))[:7]
            mes_real[kr] = mes_real.get(kr, 0) + v
    keys = sorted(set(mes_plan) | set(mes_real))
    if keys:
        _add_image(doc, chart_financeiro(finances), width_cm=16.5)
        tp = sum(mes_plan.get(k, 0) for k in keys)
        tr = sum(mes_real.get(k, 0) for k in keys)
        diff = tp - tr
        pct = (diff / tp * 100) if tp else 0
        if abs(pct) <= 5:
            parecer = f'✅ Planejamento ASSERTIVO ({pct:.1f}%)'
        elif pct > 5:
            parecer = f'⚠️ SUPERESTIMADO ({abs(pct):.1f}% abaixo do planejado)'
        else:
            parecer = f'🔴 SUBESTIMADO ({abs(pct):.1f}% acima do planejado)'
        pp = doc.add_paragraph()
        r = pp.add_run(parecer)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = FONT_NAME

        acp = acr = 0.0
        fin_rows = []
        for k in keys:
            pv, rv = mes_plan.get(k, 0), mes_real.get(k, 0)
            acp += pv
            acr += rv
            fin_rows.append([k, _fmt_brl(pv), _fmt_brl(rv), _fmt_brl(pv - rv), _fmt_brl(acp), _fmt_brl(acr)])
        fin_rows.append([('TOTAL', {'bold': True}), (_fmt_brl(tp), {'bold': True}), (_fmt_brl(tr), {'bold': True}),
                          (_fmt_brl(diff), {'bold': True}), (_fmt_brl(acp), {'bold': True}), (_fmt_brl(acr), {'bold': True})])
        _add_table(doc, ['Mês', 'Planejado', 'Realizado', 'Diferença', 'Acum. Plan.', 'Acum. Real.'],
                    fin_rows, small=True)
    else:
        doc.add_paragraph('Sem lançamentos financeiros de saída registrados.')

    # ── Riscos ────────────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, '6. Riscos', size=16)
    if risks:
        risk_rows = []
        for r in risks:
            cl = _classe_risco(r.get('probabilidade'), r.get('impacto'))
            cor = RGBColor.from_string(RISCO_CORES[cl].lstrip('#'))
            prob_i = max(1, min(5, int(r.get('probabilidade') or 1)))
            imp_i = max(1, min(5, int(r.get('impacto') or 1)))
            risk_rows.append([
                r.get('descricao', ''), r.get('categoria') or '—',
                PROB_LABELS[prob_i - 1], IMP_LABELS[imp_i - 1],
                (cl, {'bold': True, 'color': cor}),
                r.get('resposta') or '—', r.get('responsavel') or '—',
            ])
        _add_table(doc, ['Descrição', 'Categoria', 'Probabilidade', 'Impacto', 'Classificação', 'Resposta', 'Responsável'],
                    risk_rows, small=True)
    else:
        doc.add_paragraph('Nenhum risco registrado.')

    # ── Controle de Projetos ──────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, '7. Controle de Projetos (Documentos)', size=16)
    if docs:
        _add_image(doc, chart_controle_status(docs), width_cm=13)
        ctrl_rows = []
        for d in docs:
            st = DOC_STATUS.get(d['status'], {'l': d['status'] or '—', 'c': '#64748B'})
            cor = RGBColor.from_string(st['c'].lstrip('#'))
            ctrl_rows.append([
                d['codigo'] or '—', d['atividade'] or '—', d['revisao'] or '—',
                d['responsavel'] or '—', f"{d['percentual'] or 0}%",
                (st['l'], {'bold': True, 'color': cor}),
            ])
        _add_table(doc, ['Nº Doc.', 'Atividade', 'Rev.', 'Responsável', '%', 'Status'], ctrl_rows, small=True)
    else:
        doc.add_paragraph('Nenhum documento cadastrado no Controle de Projetos.')

    # ── Rodapé ────────────────────────────────────────────────────────
    doc.add_paragraph()
    fp = doc.add_paragraph(f'Relatório gerado automaticamente pelo BK ERP em {datetime.now().strftime("%d/%m/%Y %H:%M")}.')
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.runs[0].italic = True
    fp.runs[0].font.size = Pt(8.5)
    fp.runs[0].font.color.rgb = MUTED_RGB

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
